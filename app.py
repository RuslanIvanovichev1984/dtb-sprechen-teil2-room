"""
Teil-2-Raum-Prototyp: zwei Schueler sprechen jeweils in ihr EIGENES Handy
(Kanal A und B), synchron gestartet ueber diesen Server. Start kann von
JEDEM Geraet (Lehrer-Laptop ODER einem der beiden Schueler) ausgeloest
werden, sobald beide verbunden sind.

Der Server speichert NIE Audio oder Groq-Keys - jedes Handy transkribiert
lokal mit seinem eigenen Key und schickt nur Text-Schnipsel (push-to-talk
Fragmente mit eigener Zeit) an den Server. Die abschliessende Bewertung
macht irgendein Geraet, das die Seite offen hat, mit seinem EIGENEN Key
(reiner Text-Call, kein Audio) und schickt nur das Bewertungs-JSON an den
Server. Aus Transkript+Namen+Bewertung baut der Server bei Bedarf ein
herunterladbares Word-Dokument fuer die Lehrkraft.
"""
import io
import os
import random
import smtplib
import threading
import time
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

from flask import Flask, jsonify, render_template, request, send_file
from docx import Document
from docx.shared import RGBColor

GRUEN = RGBColor(0x2E, 0x7D, 0x32)
ROT = RGBColor(0xC6, 0x28, 0x28)

app = Flask(__name__)

# In-memory only - reicht fuer einen Klassenraum-Prototyp, kein Neustart-Schutz.
SESSIONS = {}
SESSION_TTL = 60 * 60  # 1 Stunde

TEACHER_EMAIL = os.environ.get("TEACHER_EMAIL", "rivanovichev@gmail.com")
GMAIL_USER = os.environ.get("GMAIL_USER", "rivanovichev@gmail.com")
GMAIL_APP_PASSWORD = os.environ.get("GMAIL_APP_PASSWORD", "")


def _new_code():
    while True:
        code = str(random.randint(100000, 999999))
        if code not in SESSIONS:
            return code


def _cleanup():
    now = time.time()
    stale = [c for c, s in SESSIONS.items() if now - s["created_at"] > SESSION_TTL]
    for c in stale:
        del SESSIONS[c]


def _fmt_time(seconds):
    m = int(seconds // 60)
    sec = int(seconds % 60)
    return f"{m:02d}:{sec:02d}"


def _merged_fragments(s):
    a_name = s["names"].get("a") or "Schueler A"
    b_name = s["names"].get("b") or "Schueler B"
    a = [{**f, "speaker": f"{a_name} (A)"} for f in s["transcripts"].get("a", {}).get("fragments", [])]
    b = [{**f, "speaker": f"{b_name} (B)"} for f in s["transcripts"].get("b", {}).get("fragments", [])]
    merged = sorted(a + b, key=lambda f: f["start"])
    return merged, a_name, b_name


@app.route("/")
def control_page():
    return render_template("index.html", role="control", preset_code="")


@app.route("/join/<code>")
def phone_page(code):
    return render_template("index.html", role="phone", preset_code=code)


@app.route("/teil1")
def teil1_page():
    return render_template("teil1.html")


@app.route("/teil3")
def teil3_page():
    return render_template("teil3.html", role="control", preset_code="")


@app.route("/teil3/join/<code>")
def teil3_phone_page(code):
    return render_template("teil3.html", role="phone", preset_code=code)


@app.route("/api/teil1/docx", methods=["POST"])
def api_teil1_docx():
    body = request.get_json(force=True)
    name = (body.get("name") or "").strip() or "Lernende-r"
    thema = body.get("thema", "")
    transkript_1a = body.get("transkript_1a", "")
    fragen_1b = body.get("fragen_1b", [])
    transkript_1b = body.get("transkript_1b", "")
    transkript_1c = body.get("transkript_1c", "")
    ev = body.get("eval", {})

    doc = Document()
    doc.add_heading("DTB B2 — Sprechen Teil 1 — Ergebnis", level=1)
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Thema: {thema}")

    doc.add_heading("1A — Monolog", level=2)
    doc.add_paragraph(transkript_1a)

    doc.add_heading("1B — Prüferfragen", level=2)
    for i, f in enumerate(fragen_1b, 1):
        doc.add_paragraph(f"{i}. {f}")
    doc.add_heading("1B — Antwort", level=3)
    doc.add_paragraph(transkript_1b)

    doc.add_heading("1C — Erläuterung in eigenen Worten", level=2)
    doc.add_paragraph(transkript_1c)

    doc.add_heading(f"Bewertung — {name}", level=2)

    for krit, label in (
        ("ki_1a", "K-I Teil 1A — Monolog"),
        ("ki_1b", "K-I Teil 1B — Prüferfragen"),
        ("ki_1c", "K-I Teil 1C — Erläuterung"),
        ("kii",   "K-II Aussprache/Intonation"),
        ("kiii",  "K-III Formale Richtigkeit"),
        ("kiv",   "K-IV Spektrum sprachlicher Mittel"),
    ):
        k = ev.get(krit, {})
        stufe = k.get("stufe", "-")
        punkte = k.get("punkte", "-")
        bestanden_krit = stufe in ("A", "B")
        icon = "✅" if bestanden_krit else "❌"
        p = doc.add_paragraph()
        r = p.add_run(f"{icon} {label}: Stufe {stufe} ({punkte} P.)")
        r.bold = True
        r.font.color.rgb = GRUEN if bestanden_krit else ROT
        doc.add_paragraph(k.get("kommentar", ""))

    gesamt = ev.get("gesamt_p", "-")
    max_p = ev.get("max_p", 24)
    p = doc.add_paragraph()
    p.add_run(f"Gesamt: {gesamt} / {max_p} Punkte").bold = True

    if ev.get("staerken"):
        doc.add_heading("Stärken", level=3)
        for st in ev["staerken"]:
            par = doc.add_paragraph(style="List Bullet")
            par.add_run(st).font.color.rgb = GRUEN

    if ev.get("verbesserungen"):
        doc.add_heading("Verbesserungsvorschläge", level=3)
        for v in ev["verbesserungen"]:
            par = doc.add_paragraph(style="List Bullet")
            par.add_run(v).font.color.rgb = ROT

    doc.add_heading("Feedback", level=3)
    doc.add_paragraph(ev.get("feedback", ""))

    muster = ev.get("musterantworten")
    if muster:
        doc.add_heading("Ideale Antworten (B2-Niveau)", level=2)
        doc.add_paragraph(
            "Die folgenden Versionen zeigen, wie die einzelnen Teile auf "
            "prüfungsausreichendem B2-Niveau hätten klingen können. "
            "Themen und Argumente wurden beibehalten, nur die sprachlichen "
            "Konstruktionen wurden verbessert."
        )
        if muster.get("monolog_1a"):
            doc.add_heading("1A — Idealer Monolog", level=3)
            doc.add_paragraph(muster["monolog_1a"])
        antworten = muster.get("antworten_1b", [])
        if antworten:
            doc.add_heading("1B — Ideale Antworten auf die Prüferfragen", level=3)
            for i, antwort in enumerate(antworten, 1):
                frage = fragen_1b[i - 1] if i - 1 < len(fragen_1b) else f"Frage {i}"
                p = doc.add_paragraph()
                p.add_run(f"Frage {i}: {frage}").bold = True
                doc.add_paragraph(antwort)
        if muster.get("erlaeuterung_1c"):
            doc.add_heading("1C — Ideale Erläuterung", level=3)
            doc.add_paragraph(muster["erlaeuterung_1c"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"DTB_B2_Teil1_{name}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/api/create", methods=["POST"])
def api_create():
    _cleanup()
    code = _new_code()
    body = request.get_json(force=True, silent=True) or {}
    SESSIONS[code] = {
        "created_at": time.time(),
        "joined": {"a": False, "b": False},
        "names": {"a": None, "b": None},
        "status": "waiting_phones",  # waiting_phones -> ready -> armed -> done
        "start_at": None,
        "duration": None,
        "transcripts": {},  # "a" / "b" -> {fragments: [{start, end, text}]}
        "eval": None,
        "impulskarten": None,
        "teil": body.get("teil", "t2"),  # "t2" oder "t3"
        "speaking": {"a": False, "b": False},  # live indicator
        "live_fragments": [],  # [{channel, start, end, text}] - for live display during recording
        "audio": {},  # channel -> (bytes, content_type), max 10 min
    }
    return jsonify({"code": code})


@app.route("/api/session/<code>", methods=["GET"])
def api_get(code):
    s = SESSIONS.get(code)
    if not s:
        return jsonify({"error": "not_found"}), 404
    return jsonify(s)


@app.route("/api/session/<code>/impulskarten", methods=["POST"])
def api_impulskarten(code):
    s = SESSIONS.get(code)
    if not s:
        return jsonify({"error": "not_found"}), 404
    if not s.get("impulskarten"):
        body = request.get_json(force=True)
        s["impulskarten"] = {"a": body.get("a", ""), "b": body.get("b", "")}
    return jsonify(s)


@app.route("/api/session/<code>/join", methods=["POST"])
def api_join(code):
    s = SESSIONS.get(code)
    if not s:
        return jsonify({"error": "not_found"}), 404
    if not s["joined"]["a"]:
        channel = "a"
    elif not s["joined"]["b"]:
        channel = "b"
    else:
        return jsonify({"error": "full"}), 409
    body = request.get_json(silent=True) or {}
    s["joined"][channel] = True
    s["names"][channel] = (body.get("name") or "").strip() or f"Schueler {channel.upper()}"
    if s["joined"]["a"] and s["joined"]["b"]:
        s["status"] = "ready"
    return jsonify({"channel": channel, "session": s})


@app.route("/api/session/<code>/start", methods=["POST"])
def api_start(code):
    s = SESSIONS.get(code)
    if not s:
        return jsonify({"error": "not_found"}), 404
    if s["status"] != "ready":
        # Schon gestartet (z.B. beide Seiten haben fast gleichzeitig gedrueckt) - ignorieren.
        return jsonify(s)
    body = request.get_json(force=True)
    duration = int(body.get("duration", 180))
    s["duration"] = duration
    s["start_at"] = time.time() + 3.0  # 3 Sek. Vorlauf, damit beide Handys rechtzeitig starten
    s["status"] = "armed"
    s["transcripts"] = {}
    s["eval"] = None
    return jsonify(s)


@app.route("/api/session/<code>/transcript/<channel>", methods=["POST"])
def api_transcript(code, channel):
    if channel not in ("a", "b"):
        return jsonify({"error": "bad_channel"}), 400
    s = SESSIONS.get(code)
    if not s:
        return jsonify({"error": "not_found"}), 404
    body = request.get_json(force=True)
    s["transcripts"][channel] = {
        "fragments": body.get("fragments", []),  # [{start, end, text}] - push-to-talk Schnipsel, eigene Geraet-Zeit
    }
    if "a" in s["transcripts"] and "b" in s["transcripts"]:
        s["status"] = "done"
    return jsonify(s)


@app.route("/api/session/<code>/speaking", methods=["POST"])
def api_speaking(code):
    s = SESSIONS.get(code)
    if not s:
        return jsonify({"error": "not_found"}), 404
    body = request.get_json(force=True)
    channel = body.get("channel")
    if channel in ("a", "b"):
        s.setdefault("speaking", {"a": False, "b": False})
        s["speaking"][channel] = bool(body.get("speaking", False))
    return jsonify({"ok": True})


@app.route("/api/session/<code>/live_fragment", methods=["POST"])
def api_live_fragment(code):
    s = SESSIONS.get(code)
    if not s:
        return jsonify({"error": "not_found"}), 404
    body = request.get_json(force=True)
    s.setdefault("live_fragments", []).append({
        "channel": body.get("channel"),
        "start": body.get("start"),
        "end": body.get("end"),
        "text": body.get("text", ""),
    })
    return jsonify({"ok": True})


@app.route("/api/session/<code>/audio/<channel>", methods=["POST"])
def api_upload_audio(code, channel):
    s = SESSIONS.get(code)
    if not s or channel not in ("a", "b"):
        return jsonify({"error": "not_found"}), 404
    ct = request.content_type or "audio/webm"
    s["audio"][channel] = (request.data, ct)
    def _expire():
        time.sleep(600)
        if code in SESSIONS:
            SESSIONS[code].get("audio", {}).pop(channel, None)
    threading.Thread(target=_expire, daemon=True).start()
    return jsonify({"ok": True})


@app.route("/api/session/<code>/audio/<channel>", methods=["GET"])
def api_get_audio(code, channel):
    s = SESSIONS.get(code)
    if not s or channel not in ("a", "b") or channel not in s.get("audio", {}):
        return jsonify({"error": "not_found"}), 404
    data, ct = s["audio"][channel]
    return send_file(io.BytesIO(data), mimetype=ct, as_attachment=False)


@app.route("/api/session/<code>/eval", methods=["POST"])
def api_eval(code):
    s = SESSIONS.get(code)
    if not s:
        return jsonify({"error": "not_found"}), 404
    if not s.get("eval"):
        s["eval"] = request.get_json(force=True)
        threading.Thread(
            target=_send_word_to_teacher, args=(code, s), daemon=True
        ).start()
    return jsonify(s)


def _build_teil2_docx(code, s):
    merged, a_name, b_name = _merged_fragments(s)
    ev = s["eval"]

    doc = Document()
    doc.add_heading("DTB B2 — Sprechen Teil 2 — Ergebnis", level=1)
    doc.add_paragraph(f"Sitzungscode: {code}")
    doc.add_paragraph(f"Teilnehmer: {a_name} & {b_name}")

    skala = doc.add_paragraph()
    skala.add_run("Bewertungsskala: ").bold = True
    skala.add_run("A = 10 P. (bestanden)  |  B = 7,5 P. (bestanden)  |  C = 4 P. (nicht bestanden)  |  D = 0 P. (nicht bestanden)")

    doc.add_heading("Dialog-Transkript", level=2)
    for frag in merged:
        doc.add_paragraph(f"[{_fmt_time(round(frag['start']))}] {frag['speaker']}: {frag['text']}")

    for key, name in (("schueler_a", a_name), ("schueler_b", b_name)):
        res = ev.get(key, {})
        doc.add_heading(f"Bewertung — {name}", level=2)
        for krit, label in (
            ("ki", "K-I Kommunikative Aufgabenbewältigung"),
            ("kii", "K-II Aussprache/Intonation"),
            ("kiii", "K-III Formale Richtigkeit"),
            ("kiv", "K-IV Spektrum sprachlicher Mittel"),
        ):
            k = res.get(krit, {})
            stufe = k.get("stufe", "-")
            bestanden_krit = stufe in ("A", "B")
            icon = "✅" if bestanden_krit else "❌"
            p = doc.add_paragraph()
            r = p.add_run(f"{icon} {label}: Stufe {stufe}")
            r.bold = True
            r.font.color.rgb = GRUEN if bestanden_krit else ROT
            doc.add_paragraph(k.get("kommentar", ""))

        bestanden = res.get("bestanden_einschaetzung")
        p = doc.add_paragraph()
        r = p.add_run("✅ ausreichend für Teil 2" if bestanden else "❌ noch nicht ausreichend für Teil 2")
        r.bold = True
        r.font.color.rgb = GRUEN if bestanden else ROT

        if res.get("staerken"):
            doc.add_heading("Stärken", level=3)
            for st in res["staerken"]:
                par = doc.add_paragraph(style="List Bullet")
                par.add_run(st).font.color.rgb = GRUEN

        if res.get("verbesserungen"):
            doc.add_heading("Verbesserungsvorschläge", level=3)
            for v in res["verbesserungen"]:
                par = doc.add_paragraph(style="List Bullet")
                par.add_run(v).font.color.rgb = ROT

        doc.add_heading("Feedback", level=3)
        doc.add_paragraph(res.get("feedback", ""))

    musterdialog = ev.get("musterdialog")
    if musterdialog:
        doc.add_heading("Ideales Muster-Gespräch (B2-Niveau)", level=2)
        doc.add_paragraph(
            "Die folgende korrigierte Version zeigt, wie das Gespräch auf "
            "prüfungsausreichendem B2-Niveau hätte klingen können. "
            "Themen und Ideen der Lernenden wurden beibehalten, nur die "
            "sprachlichen Konstruktionen wurden verbessert."
        )
        for turn in musterdialog:
            p = doc.add_paragraph()
            r = p.add_run(f"{turn.get('speaker', '')}: ")
            r.bold = True
            p.add_run(turn.get("text", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _build_teil3_docx(code, s):
    merged, a_name, b_name = _merged_fragments(s)
    ev = s["eval"]
    karte_text = (s.get("impulskarten") or {}).get("a", "")

    doc = Document()
    doc.add_heading("DTB B2 — Sprechen Teil 3 — Ergebnis", level=1)
    doc.add_paragraph(f"Sitzungscode: {code}")
    doc.add_paragraph(f"Teilnehmer: {a_name} & {b_name}")

    if karte_text:
        doc.add_heading("Situationskarte", level=2)
        doc.add_paragraph(karte_text)

    skala = doc.add_paragraph()
    skala.add_run("Bewertungsskala: ").bold = True
    skala.add_run("A = 10 P. (bestanden)  |  B = 7,5 P. (bestanden)  |  C = 4 P. (nicht bestanden)  |  D = 0 P. (nicht bestanden)")

    doc.add_heading("Dialog-Transkript", level=2)
    for frag in merged:
        doc.add_paragraph(f"[{_fmt_time(round(frag['start']))}] {frag['speaker']}: {frag['text']}")

    for key, name in (("schueler_a", a_name), ("schueler_b", b_name)):
        res = ev.get(key, {})
        doc.add_heading(f"Bewertung — {name}", level=2)
        for krit, label in (
            ("ki",   "K-I Kommunikative Aufgabenbewältigung"),
            ("kii",  "K-II Aussprache/Intonation"),
            ("kiii", "K-III Formale Richtigkeit"),
            ("kiv",  "K-IV Spektrum sprachlicher Mittel"),
        ):
            k = res.get(krit, {})
            stufe = k.get("stufe", "-")
            bestanden_krit = stufe in ("A", "B")
            icon = "✅" if bestanden_krit else "❌"
            p = doc.add_paragraph()
            r = p.add_run(f"{icon} {label}: Stufe {stufe}")
            r.bold = True
            r.font.color.rgb = GRUEN if bestanden_krit else ROT
            doc.add_paragraph(k.get("kommentar", ""))

        aufgabenvert = res.get("aufgabenverteilung")
        p = doc.add_paragraph()
        r = p.add_run(
            "✅ Aufgabenverteilung vorhanden" if aufgabenvert
            else "❌ Aufgabenverteilung fehlt (K-I max. Stufe C)"
        )
        r.bold = True
        r.font.color.rgb = GRUEN if aufgabenvert else ROT

        bestanden = res.get("bestanden_einschaetzung")
        p = doc.add_paragraph()
        r = p.add_run("✅ ausreichend für Teil 3" if bestanden else "❌ noch nicht ausreichend für Teil 3")
        r.bold = True
        r.font.color.rgb = GRUEN if bestanden else ROT

        if res.get("staerken"):
            doc.add_heading("Stärken", level=3)
            for st in res["staerken"]:
                par = doc.add_paragraph(style="List Bullet")
                par.add_run(st).font.color.rgb = GRUEN

        if res.get("verbesserungen"):
            doc.add_heading("Verbesserungsvorschläge", level=3)
            for v in res["verbesserungen"]:
                par = doc.add_paragraph(style="List Bullet")
                par.add_run(v).font.color.rgb = ROT

        doc.add_heading("Feedback", level=3)
        doc.add_paragraph(res.get("feedback", ""))

    musterdialog = ev.get("musterdialog")
    if musterdialog:
        doc.add_heading("Ideales Muster-Gespräch (B2-Niveau)", level=2)
        doc.add_paragraph(
            "Die folgende korrigierte Version zeigt, wie die Diskussion auf "
            "prüfungsausreichendem B2-Niveau ausgesehen hätte — mit Lösungsvorschlägen, "
            "Reaktionen und Aufgabenverteilung am Ende."
        )
        for turn in musterdialog:
            p = doc.add_paragraph()
            r = p.add_run(f"{turn.get('speaker', '')}: ")
            r.bold = True
            p.add_run(turn.get("text", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _send_word_to_teacher(code, s):
    if not GMAIL_APP_PASSWORD:
        return
    try:
        teil = s.get("teil", "t2")
        buf = _build_teil3_docx(code, s) if teil == "t3" else _build_teil2_docx(code, s)
        a_name = s["names"].get("a") or "Schueler A"
        b_name = s["names"].get("b") or "Schueler B"

        msg = MIMEMultipart()
        msg["From"] = GMAIL_USER
        msg["To"] = TEACHER_EMAIL
        teil_nr = "3" if s.get("teil") == "t3" else "2"
        msg["Subject"] = f"DTB B2 Teil {teil_nr} — {a_name} & {b_name} (Sitzung {code})"
        msg.attach(MIMEText(
            f"Neue Sitzung abgeschlossen.\n\nTeil: {teil_nr}\nCode: {code}\nTeilnehmer: {a_name} & {b_name}",
            "plain", "utf-8"
        ))
        part = MIMEBase("application", "octet-stream")
        part.set_payload(buf.read())
        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            f'attachment; filename="DTB_B2_Teil2_{code}.docx"'
        )
        msg.attach(part)

        with smtplib.SMTP("smtp.gmail.com", 587) as srv:
            srv.ehlo()
            srv.starttls()
            srv.login(GMAIL_USER, GMAIL_APP_PASSWORD)
            srv.send_message(msg)
        print(f"[email] Word gesendet fuer Sitzung {code}")
    except Exception as e:
        print(f"[email] Fehler beim Senden: {e}")


@app.route("/api/test_email")
def api_test_email():
    if not GMAIL_APP_PASSWORD:
        return jsonify({"error": "GMAIL_APP_PASSWORD not set"}), 500
    try:
        msg = MIMEMultipart()
        msg["From"] = GMAIL_USER
        msg["To"] = TEACHER_EMAIL
        msg["Subject"] = "DTB Sprechen App — Email-Test"
        msg.attach(MIMEText("Email-Versand funktioniert.", "plain", "utf-8"))
        with smtplib.SMTP("smtp.gmail.com", 587) as srv:
            srv.ehlo()
            srv.starttls()
            srv.login(GMAIL_USER, GMAIL_APP_PASSWORD)
            srv.send_message(msg)
        return jsonify({"ok": True, "to": TEACHER_EMAIL})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/session/<code>/docx")
def api_docx(code):
    s = SESSIONS.get(code)
    if not s or not s.get("eval"):
        return jsonify({"error": "not_ready"}), 404
    teil = s.get("teil", "t2")
    if teil == "t3":
        buf = _build_teil3_docx(code, s)
        filename = f"DTB_B2_Teil3_{code}.docx"
    else:
        buf = _build_teil2_docx(code, s)
        filename = f"DTB_B2_Teil2_{code}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
